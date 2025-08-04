import io
import os
from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentProcessor:
    def __init__(self):
        pass

    def process_pdf(self, pdf_file):
        """Extract text from a PDF file object or path"""
        text = ""

        if isinstance(pdf_file, str):
            with open(pdf_file, "rb") as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        else:
            # File object (like from Flask)
            pdf_file.seek(0)
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text.strip()

    def process_word(self, word_file):
        """Extract text from a Word (.docx) file object or path"""
        text = ""

        if isinstance(word_file, str):
            doc = DocxDocument(word_file)
        else:
            word_file.seek(0)
            doc = DocxDocument(io.BytesIO(word_file.read()))

        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"

        return text.strip()

    def extract_relevant_information(self, text, query):
        """
        Placeholder: Add NLP/LLM-based logic here to retrieve matching clauses
        from `text` based on the user `query`.
        """
        # Example dummy clause extraction
        if "knee surgery" in query.lower() and "covered" in text.lower():
            clause = "Knee surgery is covered after 3 months of policy duration."
            return {"clause": clause}
        else:
            return {"clause": "No relevant clause found."}

    def evaluate_decision(self, extracted_info, query):
        """
        Evaluate whether the query meets criteria based on extracted clauses.
        Real logic should use embeddings, rule matching, LLMs, etc.
        """
        clause = extracted_info.get("clause", "")
        if "covered" in clause.lower():
            return {
                "decision": "Approved",
                "amount": "As per policy terms",
                "justification": clause
            }
        else:
            return {
                "decision": "Rejected",
                "amount": None,
                "justification": "No matching policy clause found for the query"
            }

    def generate_response(self, decision_data):
        """
        Format the final structured response
        """
        return {
            "decision": decision_data.get("decision"),
            "amount": decision_data.get("amount"),
            "justification": decision_data.get("justification")
        }
