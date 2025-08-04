import io
import os
from typing import Union, Optional

from pypdf import PdfReader
from docx import Document
from email import policy
from email.parser import BytesParser


class DocumentService:
    SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "eml"}

    def __init__(self):
        pass

    def extract_text(self, file: Union[str, io.BytesIO]) -> str:
        """Extracts text from a supported document type (PDF, DOCX, TXT, EML)."""

        ext = self._get_file_extension(file)

        if ext == "pdf":
            return self._extract_text_from_pdf(file)
        elif ext == "docx":
            return self._extract_text_from_docx(file)
        elif ext == "txt":
            return self._extract_text_from_txt(file)
        elif ext == "eml":
            return self._extract_text_from_eml(file)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def extract_metadata(self, file: Union[str, io.BytesIO]) -> Optional[dict]:
        """Extracts metadata from a PDF file only."""

        ext = self._get_file_extension(file)
        if ext != "pdf":
            raise ValueError("Metadata extraction is only supported for PDF files.")

        if isinstance(file, str):
            with open(file, "rb") as f:
                reader = PdfReader(f)
                return dict(reader.metadata or {})
        else:
            file.seek(0)
            reader = PdfReader(file)
            return dict(reader.metadata or {})

    def process_document(self, file: Union[str, io.BytesIO]) -> dict:
        """Returns both extracted text and metadata (for PDFs)."""

        text = self.extract_text(file)
        ext = self._get_file_extension(file)

        metadata = None
        if ext == "pdf":
            metadata = self.extract_metadata(file)

        return {
            "text": text,
            "metadata": metadata
        }

    def _get_file_extension(self, file: Union[str, io.BytesIO]) -> str:
        """Safely extract file extension (e.g. 'pdf', 'docx') from file or stream."""

        if isinstance(file, str):
            ext = os.path.splitext(file)[1].lower().lstrip(".")
        else:
            filename = getattr(file, 'filename', '')
            ext = os.path.splitext(filename)[1].lower().lstrip(".")
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}")
        return ext

    def _extract_text_from_pdf(self, file: Union[str, io.BytesIO]) -> str:
        text = ""
        if isinstance(file, str):
            with open(file, "rb") as f:
                reader = PdfReader(f)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
        else:
            file.seek(0)
            reader = PdfReader(file)
            text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return text.strip()

    def _extract_text_from_docx(self, file: Union[str, io.BytesIO]) -> str:
        if isinstance(file, str):
            doc = Document(file)
        else:
            file.seek(0)
            doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs]).strip()

    def _extract_text_from_txt(self, file: Union[str, io.BytesIO]) -> str:
        if isinstance(file, str):
            with open(file, 'r', encoding='utf-8') as f:
                return f.read().strip()
        else:
            file.seek(0)
            return file.read().decode('utf-8').strip()

    def _extract_text_from_eml(self, file: Union[str, io.BytesIO]) -> str:
        if isinstance(file, str):
            with open(file, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
        else:
            file.seek(0)
            msg = BytesParser(policy=policy.default).parse(file)
        return msg.get_body(preferencelist=('plain')).get_content().strip()
